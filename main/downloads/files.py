import docx
from docx.shared import Pt
import win32com.client
import webbrowser
from docx.shared import Inches
import win32com.client
def download_reports(p_name,d_name,s_spes,date,time,pres,report):
    url='file:///C:/programming/Hackathontyl/'
    g_path='C:/Program Files/Google/Chrome/Application/chrome.exe %s'
    # C:\Program Files\Google\Chrome\Application
    def billing(name, email, product, unit, price):

        document = docx.Document()
        
        document.add_heading('Invoice', 0)
        p1 = document.add_heading('Dear ',1)
        p1.add_run(name).bold=True
        p1.add_run(',')

        p21 = document.add_paragraph("  ")
        p21.add_run('GST Registration No : ').bold = True
        p21.add_run('wer-ytre-65432345')

        p21 = document.add_paragraph(" ")
        
        p21.add_run('PAN No : ').bold = True
        p21.add_run('AVSBQH267J')

        p21 = document.add_paragraph(" ")

        p21.add_run('Invoice Details : ').bold = True
        p21.add_run('SCCH-273892')

        p21 = document.add_paragraph(" ")
        
        p2 = document.add_heading('Please find attached invoice for your recent purchase of Medicine ',1)

        p2.add_run(str(unit)).bold = True
        p2.add_run(' units of ')
        p2.add_run(product).bold=True
        p2.add_run('.')

        [document.add_paragraph('') for _ in range(1)]
        
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product Name'
        hdr_cells[1].text = 'Units'
        hdr_cells[2].text = 'Unit Price'
        hdr_cells[3].text = 'Total Price'
        for i in range(4):
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        row_cells = table.add_row().cells
        row_cells[0].text = product
        row_cells[1].text = f'{unit:,.2f}'
        row_cells[2].text = f'{price:,.2f}'
        row_cells[3].text = f'{unit * price:,.2f}'
        
        [document.add_paragraph('') for _ in range(10)]

        document.add_paragraph('''*iMediCare Seller Services Pvt. Ltd., iMediCare Retail India Pvt. Ltd. (only where iMediCare Retail India Pvt. Ltd. fulfillment center is co-located) 
    Customers desirous of availing input GST credit are requested to create a Business account and purchase on iMediCare.in/business from Business eligible offers Please note that this invoice is not a demand for payment*''')

        document.save(name+'.docx')



    def report(name,dname,spec,text):
        document = docx.Document()
        document.add_heading(r'Patient Report', 0)

        p0 = document.add_heading('Details  :',1)

        p1 = document.add_heading('Patient Name  :  ',2)
        p1.add_run(name)

        p2 = document.add_heading('Doctor Name  :  ',2)
        p2.add_run(dname)
        
        p3 = document.add_heading('Treatment For  :  ',2)
        p3.add_run(spec)

        document.add_heading('Description  :',1)
        p4 = document.add_heading(text,2)

        name=name.replace(' ','_')+'_report'
        document.save(name+'.docx')
        webbrowser.get(g_path).open(url+name+'.docx')
        


    def precipitation(name,text):
        document = docx.Document()
        document.add_heading(r'Doctor Precipitation  ', 0)
        document.add_heading(text,1)
        
        document.add_heading(text,2)

        name=name.replace(' ','_')+'_precipitation'
        document.save(name+'.docx')
        webbrowser.get(g_path).open(url+name+'.docx')


    billing("rakshiht","Rakshith@gail.com","cosmic",2,40)

    precipitation(p_name,pres)
    report("Kiran Babu","Dr.Vishal","Dentist","""The brain is the control center of the body. It controls thoughts, memory, speech, and movement. It regulates the function of many organs. When the brain is healthy, it works quickly and automatically. However, when problems occur, the results can be devastating.

    Inflammation in the brain can lead to problems such as vision loss, weakness and paralysis. Loss of brain cells, which happens if you suffer a stroke, can affect your ability to think clearly. Brain tumors can also press on nerves and affect brain function. Some brain diseases are genetic. And we do not know what causes some brain diseases, such as Alzheimer's disease.

    The symptoms of brain diseases vary widely depending on the specific problem. In some cases, damage is permanent. In other cases, treatments such as surgery, medicines, or physical therapy can correct the source of the problem or improve symptoms.""")